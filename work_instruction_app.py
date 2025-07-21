import streamlit as st
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from fpdf import FPDF
from io import BytesIO
from datetime import datetime
import os

# Constants
LOGO_PATH = "logo.jpg"  # logo must be in the same directory
COMPANY_NAME = "BOROSIL RENEWBALES LIMITED"
DEFAULT_CLAUSES = [
    "Scope",
    "Purpose",
    "Frequency",
    "Sample size",
    "Resources required",
    "Responsibility",
    "Procedure steps",
    "PPEs matrix",
    "EHS Requirements",
    "Reference documents",
    "Revision history"
]

st.set_page_config(page_title="Work Instruction Generator")
st.title("📝 Work Instruction Generator v1")
st.markdown("""
    <style>
    /* Aggressively hide all file uploader info text except label and button */
    div[data-testid="stFileUploader"] *:not(label):not(button) {
        visibility: hidden !important;
        height: 0 !important;
        margin: 0 !important;
        padding: 0 !important;
    }
    div[data-testid="stFileUploader"] label {
        font-size: 14px !important;
        visibility: visible !important;
    }
    div[data-testid="stFileUploader"] button {
        visibility: visible !important;
    }
    </style>
""", unsafe_allow_html=True)

# Input fields
wi_title = st.text_input("Title of Work Instruction")

# Department to document prefix mapping
dept_prefix_map = {
    "Quality": "QA/L3/",
    "Batch House": "BH/L3/",
    "Furnace": "FUR/L3/",
    "Rolling machine": "ROM/L3/",
    "Lehr & cutting": "LRC/L3/",
    "Annealed Packing": "PRD/L3/",
    "Grinding & Drilling": "GRND/L3/",
    "Grid Printing": "GRID/L3/",
    "ARC": "ARC/L3/",
    "Tempering": "TEMP/L3/",
    "Final packing": "PACK/L3/",
    "Warehouse": "WH/L3/",
    "Box Yard": "BY/L3/",
    "Lab": "LAB/L3/",
    "Mechanical": "MECH/L3/",
    "Electrial": "ELE/L3/",
    "Instrumentation": "INST/L3/",
    "Utility": "UTIL/L3/",
    "EHS": "EHS/L3/",
    "HR": "HR/L3/",
    "Admin": "ADMIN/L3/",
    "Purchase": "PUR/L3/",
    "IT": "IT/L3/",
    "Marketing": "MKT/L3/",
    "MR": "MR/L3/"
}

department = st.selectbox("Department", list(dept_prefix_map.keys()))

# Use session state to update doc_no prefix when department changes
if "last_department" not in st.session_state:
    st.session_state.last_department = department
if "doc_no" not in st.session_state:
    st.session_state.doc_no = dept_prefix_map.get(department, "") + "001"
if department != st.session_state.last_department:
    st.session_state.doc_no = dept_prefix_map.get(department, "") + "001"
    st.session_state.last_department = department
doc_no = st.text_input("Document No:", st.session_state.doc_no, key="doc_no")
issue_date = st.date_input("Issue Date", datetime.today())
rev_no = st.text_input("Revision No:", "00")
rev_date = st.date_input("Revision Date", datetime.today())


clauses = []
st.subheader("Clauses")
for i, clause in enumerate(DEFAULT_CLAUSES):
    if clause == "Resources required":
        st.markdown(f"**{i+1}. {clause}**")
        machine = st.text_area(f"{i+1}.1 Machine", key="machine")
        material = st.text_area(f"{i+1}.2 Material", key="material")
        man = st.text_area(f"{i+1}.3 Man", key="man")
        clauses.append((clause, {"machine": machine, "material": material, "man": man}))
    elif clause == "PPEs matrix":
        st.markdown(f"**{i+1}. {clause}**")
        text = st.text_area(f"{i+1}. {clause}", key=f"clause_{i}")
        # PPEs options and images
        ppe_options = [
            {"name": "Goggle", "image": "goggle.png"},
            {"name": "Shoe", "image": "shoe.png"},
            {"name": "Helmet", "image": "helmet.png"},
            {"name": "Gloves", "image": "gloves.png"},
            {"name": "Mask", "image": "mask.png"},
            {"name": "Apron", "image": "apron.png"}
        ]
        ppe_selected = st.multiselect("Select PPEs for PPEs Matrix", [ppe["name"] for ppe in ppe_options])
        clauses.append((clause, text))
    elif clause == "Procedure steps":
        st.markdown(f"**{i+1}. {clause}**")
        num_steps = st.number_input("Number of steps", min_value=1, max_value=20, value=1, key="num_steps")
        steps = []
        for step_idx in range(int(num_steps)):
            cols = st.columns([3,1,1])
            detail = cols[0].text_area(f"Step {step_idx+1} Detail", key=f"step_detail_{step_idx}")
            with cols[1]:
                st.markdown('<div style="text-align:center;margin-bottom:4px;font-size:16px;">Attach pic 1</div>', unsafe_allow_html=True)
                img1 = st.file_uploader("Attach pic.", type=["png", "jpg", "jpeg"], key=f"step_img1_{step_idx}", label_visibility="collapsed")
            with cols[2]:
                st.markdown('<div style="text-align:center;margin-bottom:4px;font-size:16px;">Attach pic 2</div>', unsafe_allow_html=True)
                img2 = st.file_uploader("Attach pic.", type=["png", "jpg", "jpeg"], key=f"step_img2_{step_idx}", label_visibility="collapsed")
            steps.append({"detail": detail, "images": [img1, img2]})
        clauses.append((clause, steps))
    else:
        st.markdown(f"**{i+1}. {clause}**")
        text = st.text_area(f"{i+1}. {clause}", key=f"clause_{i}")
        clauses.append((clause, text))

extra_clauses = st.number_input("Add Extra Clauses", min_value=0, max_value=10, step=1)
for i in range(extra_clauses):
    title = st.text_input(f"Extra Clause {i+1} Title", key=f"extra_title_{i}")
    body = st.text_area(f"Extra Clause {i+1} Content", key=f"extra_body_{i}")
    if title:
        clauses.append((title, body))

# Footer details
prep_by = st.text_input("Prepared By")
review_by = st.text_input("Reviewed By")
approve_by = st.text_input("Approved By")


def create_docx():
    doc = Document()

    # Set page margins (1.5 cm = 0.59 inches)
    section = doc.sections[0]
    section.top_margin = Inches(0.59)
    section.bottom_margin = Inches(0.59)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.59)
    
    header = section.header
    # Remove any extra empty paragraph in header before table
    if header.paragraphs and not header.paragraphs[0].text.strip():
        p = header.paragraphs[0]._element
        p.getparent().remove(p)
    footer = section.footer

    # Header table (adjusted for margins)
    # Total width = 8.27 (A4 width) - 0.59*2 (margins) = 7.09 inches
    table = header.add_table(rows=4, cols=3, width=Inches(7.09))
    table.style = 'Table Grid'
    table.autofit = False
    # Adjust column widths proportionally
    widths = [Inches(1.5), Inches(3.59), Inches(2.2)]
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = widths[idx]

    # Merge cells for logo
    logo_cell = table.cell(0, 0)
    for i in range(1, 4):
        logo_cell.merge(table.cell(i, 0))
    paragraph = logo_cell.paragraphs[0]
    run = paragraph.add_run()
    if os.path.exists(LOGO_PATH):
        run.add_picture(LOGO_PATH, width=Inches(1.2))

    # Merge cells for company name
    name_cell = table.cell(0, 1)
    name_cell.merge(table.cell(1, 1))
    paragraph = name_cell.paragraphs[0]
    run = paragraph.add_run(COMPANY_NAME)
    run.font.size = Pt(14)  # Larger font for company name
    run.font.bold = True
    run.font.name = 'Arial'
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Merge cells for SOP title
    sop_cell = table.cell(2, 1)
    sop_cell.merge(table.cell(3, 1))
    paragraph = sop_cell.paragraphs[0]
    run = paragraph.add_run(f"WORK INSTRUCTION – {wi_title.upper()}")
    run.font.size = Pt(12)  # Medium font for WI title
    run.font.bold = True
    run.font.name = 'Arial'
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Column 3 with smaller font
    for row, text in enumerate([
        f"DOC. NO: {doc_no}",
        f"ISSUE NO. / DATE: {rev_no} / {issue_date}",
        f"REV. NO: {rev_no}",
        f"REV. DATE: {rev_date}"
    ]):
        cell = table.cell(row, 2)
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(text)
        run.font.size = Pt(9)  # Smaller font for document details
        run.font.name = 'Arial'
        run.font.bold = False

    # Vertically center align all header table cells
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            vAlign = OxmlElement('w:vAlign')
            vAlign.set(qn('w:val'), 'center')
            tcPr.append(vAlign)

    # Increase height of all header rows
    for row in table.rows:
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        trHeight = OxmlElement('w:trHeight')
        trHeight.set(qn('w:val'), '300')  # 500 twips ~0.35 cm
        trHeight.set(qn('w:hRule'), 'exact')
        trPr.append(trHeight)


    # WI title
    heading = doc.add_paragraph(wi_title)
    heading.style = doc.styles['Heading 1']
    # Set font for heading explicitly
    for run in heading.runs:
        run.font.name = 'Calibri'
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')

    dept_para = doc.add_paragraph(f"Department: {department}")
    for run in dept_para.runs:
        run.font.name = 'Calibri'
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')

    # Clause entries with manual numbering
    for idx, (title, content) in enumerate(clauses, 1):
        if title == "Resources required":
            p = doc.add_paragraph()
            run = p.add_run(f"{idx}. {title}:")
            run.bold = True
            run.font.name = 'Calibri'
            r = run._element
            r.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
            res_table = doc.add_table(rows=2, cols=3)
            res_table.style = 'Table Grid'
            res_table.autofit = True
            headers = ["Machine", "Material", "Man"]
            for col, h in enumerate(headers):
                cell = res_table.cell(0, col)
                for paragraph in cell.paragraphs:
                    run = paragraph.add_run(h)
                    run.bold = True
                    run.font.name = 'Calibri'
                    r = run._element
                    r.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph.paragraph_format.space_before = Pt(3)
                    paragraph.paragraph_format.space_after = Pt(3)
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                vAlign = OxmlElement('w:vAlign')
                vAlign.set(qn('w:val'), 'center')
                tcPr.append(vAlign)
                shd = OxmlElement('w:shd')
                shd.set(qn('w:fill'), 'B7DEE8')
                tcPr.append(shd)
            entries = [content["machine"], content["material"], content["man"]]
            for col, entry in enumerate(entries):
                cell = res_table.cell(1, col)
                cell.text = entry
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Calibri'
                        r = run._element
                        r.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph.paragraph_format.space_before = Pt(6)
                    paragraph.paragraph_format.space_after = Pt(6)
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                vAlign = OxmlElement('w:vAlign')
                vAlign.set(qn('w:val'), 'center')
                tcPr.append(vAlign)
        elif title == "Procedure steps":
            p = doc.add_paragraph()
            run = p.add_run(f"{idx}. {title}:")
            run.bold = True
            run.font.name = 'Calibri'
            r = run._element
            r.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
            steps = content
            proc_table = doc.add_table(rows=1 + len(steps), cols=3)
            proc_table.style = 'Table Grid'
            proc_table.autofit = False
            col_widths = [Inches(1/2.54), Inches(9.5/2.54), Inches(8.05/2.54)]
            for row in proc_table.rows:
                for col, cell in enumerate(row.cells):
                    cell.width = col_widths[col]
            header_cells = proc_table.rows[0].cells
            header_cells[0].text = "#"
            header_cells[1].text = "Detail"
            header_cells[2].text = "Picture"
            for col in range(3):
                for paragraph in header_cells[col].paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Calibri'
                        r = run._element
                        r.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                    run.bold = True
                    run.font.name = 'Calibri'
                    r = run._element
                    r.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
                tc = header_cells[col]._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:fill'), 'B7DEE8')
                tcPr.append(shd)
            for step_idx, step in enumerate(steps):
                row_cells = proc_table.rows[step_idx+1].cells
                row_cells[0].text = str(step_idx+1)
                for paragraph in row_cells[0].paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Calibri'
                        r = run._element
                        r.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                row_cells[1].text = step["detail"]
                for paragraph in row_cells[1].paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Calibri'
                        r = run._element
                        r.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                pic_cell = row_cells[2]
                for para in pic_cell.paragraphs:
                    para.clear()
                for img in step["images"]:
                    if img is not None:
                        try:
                            pic_cell.add_paragraph().add_run().add_picture(img, width=Inches(1.2))
                        except Exception:
                            pass
                for paragraph in pic_cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Calibri'
                        r = run._element
                        r.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif idx == 8:  # PPEs matrix clause
            p = doc.add_paragraph()
            run = p.add_run(f"{idx}. {title}:")
            run.bold = True
            run.font.name = 'Calibri'
            r = run._element
            r.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
            selected_ppe_objs = [ppe for ppe in ppe_options if ppe["name"] in ppe_selected]
            if selected_ppe_objs:
                ppe_table = doc.add_table(rows=2, cols=len(selected_ppe_objs))
                ppe_table.autofit = True
                ppe_table.style = 'Table Grid'
                for col in range(len(selected_ppe_objs)):
                    for row in range(2):
                        cell = ppe_table.cell(row, col)
                        cell.width = Inches(7.09 / max(1, len(selected_ppe_objs)))
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.name = 'Calibri'
                                r = run._element
                                r.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        tc = cell._tc
                        tcPr = tc.get_or_add_tcPr()
                        vAlign = OxmlElement('w:vAlign')
                        vAlign.set(qn('w:val'), 'center')
                        tcPr.append(vAlign)
                for col, ppe in enumerate(selected_ppe_objs):
                    cell = ppe_table.cell(0, col)
                    if os.path.exists(ppe["image"]):
                        cell.paragraphs[0].add_run().add_picture(ppe["image"], width=Inches(1))
                for col, ppe in enumerate(selected_ppe_objs):
                    cell = ppe_table.cell(1, col)
                    run = cell.paragraphs[0].add_run(ppe["name"])
                    run.font.name = 'Calibri'
                    r = run._element
                    r.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
            else:
                doc.add_paragraph("No PPEs selected.")
        else:
            p = doc.add_paragraph()
            run = p.add_run(f"{idx}. {title}:")
            run.bold = True
            run.font.name = 'Calibri'
            r = run._element
            r.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
            para = doc.add_paragraph(content)
            for run in para.runs:
                run.font.name = 'Calibri'
                r = run._element
                r.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')

    # Add approval table to footer section (3 rows, 4 columns, last column merged)
    footer_table = footer.add_table(rows=3, cols=4, width=Inches(7.09))
    footer_table.style = 'Table Grid'
    # Set column widths
    footer_widths = [Inches(2.0), Inches(2.0), Inches(2.0), Inches(1.3)]
    for row in footer_table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = footer_widths[idx]

    # Merge last column for all rows (signature)
    signature_cell = footer_table.cell(0, 3)
    for i in range(1, 3):
        signature_cell.merge(footer_table.cell(i, 3))

    # Center align all cells vertically and horizontally
    for row in footer_table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            vAlign = OxmlElement('w:vAlign')
            vAlign.set(qn('w:val'), 'center')
            tcPr.append(vAlign)

    # Increase height of second and third row
    for row_idx in [1, 2]:
        tr = footer_table.rows[row_idx]._tr
        trPr = tr.get_or_add_trPr()
        trHeight = OxmlElement('w:trHeight')
        trHeight.set(qn('w:val'), '500')  # 500 twips ~0.35 cm
        trHeight.set(qn('w:hRule'), 'exact')
        trPr.append(trHeight)

    # First row: labels
    for col, label in enumerate(["Prepared By", "Reviewed By", "Approved By"]):
        cell = footer_table.cell(0, col)
        cell.text = label
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Only set 'Signature' in the first row's merged cell
    signature_paragraph = footer_table.cell(0, 3).paragraphs[0]
    signature_paragraph.clear()
    signature_paragraph.add_run("Signature")
    signature_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Second row: names from input
    for col, name in enumerate([prep_by, review_by, approve_by]):
        cell = footer_table.cell(1, col)
        cell.text = name
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Leave merged cell in second row blank
    blank_paragraph = footer_table.cell(1, 3).paragraphs[0]
    blank_paragraph.clear()
    blank_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Third row: Page x of x using Word fields
    for col in range(3):
        footer_table.cell(2, col).text = ""
    # Insert Word field for page numbering only in merged cell
    page_paragraph = footer_table.cell(2, 3).paragraphs[0]
    page_paragraph.clear()
    run = page_paragraph.add_run("Page ")
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText1 = OxmlElement('w:instrText')
    instrText1.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText1)
    run._r.append(fldChar2)
    run._r.append(fldChar3)
    run = page_paragraph.add_run(" of ")
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText1 = OxmlElement('w:instrText')
    instrText1.text = 'NUMPAGES'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText1)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

    return doc


def generate_download(doc):
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream


if st.button("Generate Work Instruction"):
    document = create_docx()
    docx_file = generate_download(document)
    st.download_button("Download DOCX", docx_file, file_name="work_instruction.docx")
